from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

AZUL  = RGBColor(30, 58, 138)
VERDE = RGBColor(13, 148, 136)
NEGRO = RGBColor(0, 0, 0)

def font(run, bold=False, size=11, color=NEGRO):
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color

def heading_center(text, size=14, color=AZUL):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    font(r, bold=True, size=size, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p

def heading_left(text, size=11, color=AZUL):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    font(r, bold=True, size=size, color=color)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(text, indent=0.5):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_after  = Pt(4)
    for r in p.runs:
        font(r, size=10.5)
    return p

def field_row(label, value="_______________________________"):
    p = doc.add_paragraph()
    r1 = p.add_run(label + ":  ")
    font(r1, bold=True, size=10.5)
    r2 = p.add_run(value)
    font(r2, size=10.5)
    r2.underline = True
    p.paragraph_format.space_after = Pt(3)

def clause(title, text):
    heading_left(title)
    body(text)

# ── ENCABEZADO
heading_center("VERONA RIDE — VENEZUELA", 14, AZUL)
heading_center("CONTRATO DE PRESTACION DE SERVICIOS COMO CONDUCTOR INDEPENDIENTE", 13, VERDE)
doc.add_paragraph()

# ── DATOS DEL CONDUCTOR
heading_left("DATOS DEL CONDUCTOR")
field_row("Nombre completo")
field_row("Cedula de Identidad")
field_row("Telefono")
field_row("Correo electronico")
field_row("Estado / Ciudad")
doc.add_paragraph()

# ── DATOS DEL VEHICULO
heading_left("DATOS DEL VEHICULO")
field_row("Tipo de vehiculo")
field_row("Marca")
field_row("Modelo")
field_row("Ano")
field_row("Color")
field_row("Placa")
field_row("Compania de seguro")
field_row("N de poliza")
field_row("Vencimiento del seguro")
doc.add_paragraph()

# ── CLAUSULAS
heading_center("CLAUSULAS DEL CONTRATO", 12, AZUL)

clause("PRIMERA - NATURALEZA DEL SERVICIO",
    "EL CONDUCTOR prestara servicios de transporte de personas y/o encomiendas a traves de la "
    "plataforma tecnologica VERONA Ride, operando como trabajador independiente. No existe relacion "
    "laboral, de dependencia ni subordinacion entre EL CONDUCTOR y LA PLATAFORMA.")

clause("SEGUNDA - MEMBRESIA SEMANAL",
    "EL CONDUCTOR se compromete a pagar la membresia semanal correspondiente al tipo de vehiculo "
    "registrado, conforme a las siguientes tarifas vigentes:\n"
    "   Motocicleta:   USD 15,00 por semana\n"
    "   Sedan:          USD 25,00 por semana\n"
    "   SUV:             USD 30,00 por semana\n"
    "   Pick-Up:        USD 35,00 por semana\n"
    "   Plataforma:    USD 40,00 por semana\n"
    "El periodo de membresia es de viernes a jueves. El pago debe realizarse antes de iniciar "
    "operaciones cada semana.")

clause("TERCERA - COMISION",
    "VERONA Ride Venezuela opera con CERO POR CIENTO (0%) de comision. EL CONDUCTOR conserva "
    "el 100% del valor cobrado por cada servicio prestado a los usuarios de la plataforma.")

clause("CUARTA - TARIFAS Y FORMA DE COBRO",
    "El cobro de los servicios se realiza en dolares estadounidenses (USD) en efectivo. Las tarifas "
    "son calculadas automaticamente por la plataforma segun distancia recorrida, tiempo de viaje y "
    "tipo de servicio solicitado. Los montos en bolivares se expresan de manera referencial conforme "
    "a la tasa oficial del Banco Central de Venezuela (BCV) vigente al momento del servicio.")

clause("QUINTA - PENALIZACION POR CANCELACION",
    "Las cancelaciones realizadas por el pasajero despues de los 2 (dos) minutos de gracia contados "
    "desde la asignacion del conductor generan los siguientes cargos:\n"
    "   Motocicleta:                     USD 0,75 (monto fijo)\n"
    "   Sedan:                              USD 0,80 (monto fijo)\n"
    "   SUV:                                USD 0,90 (monto fijo)\n"
    "   Encomienda:                      25% de la tarifa estimada\n"
    "   Pick-Up / Plataforma / Carga:  35% de la tarifa estimada")

clause("SEXTA - CODIGO DE CONDUCTA",
    "EL CONDUCTOR se compromete a:\n"
    "   a) Mantener una calificacion promedio minima de 4,5 estrellas en la plataforma.\n"
    "   b) No operar el vehiculo bajo efectos de alcohol, drogas u otras sustancias psicoactivas.\n"
    "   c) Mantener el vehiculo en condiciones optimas de limpieza, higiene y mantenimiento mecanico.\n"
    "   d) Tratar a todos los pasajeros con respeto, amabilidad y profesionalismo en todo momento.\n"
    "   e) No fumar dentro del vehiculo durante servicios activos ni mientras espera al pasajero.\n"
    "   f) No compartir ni divulgar informacion personal de los pasajeros a terceros.\n"
    "   g) Completar los cursos de certificacion requeridos por LA PLATAFORMA en los plazos establecidos.\n"
    "   h) Presentarse con vestimenta limpia y apropiada al momento de cada servicio.\n"
    "   i) No realizar desvios de ruta sin previo consentimiento del pasajero.")

clause("SEPTIMA - SUSPENSION Y CANCELACION DE CUENTA",
    "LA PLATAFORMA se reserva el derecho de suspender temporalmente o cancelar de forma definitiva "
    "la cuenta de EL CONDUCTOR en caso de:\n"
    "   - Incumplimiento de cualquiera de las clausulas del presente contrato.\n"
    "   - Calificacion promedio sostenida inferior a 4,5 estrellas por mas de 7 dias consecutivos.\n"
    "   - Uso fraudulento de la plataforma o manipulacion de tarifas y calificaciones.\n"
    "   - Comportamiento inapropiado, agresivo o irrespetuoso reportado por uno o mas pasajeros.\n"
    "   - Conduccion bajo efectos de alcohol u otras sustancias (tolerancia cero, cancelacion inmediata).\n"
    "   - Incumplimiento en el pago de la membresia semanal.")

clause("OCTAVA - RESPONSABILIDAD DEL CONDUCTOR",
    "EL CONDUCTOR es responsable de mantener vigentes y en regla todos sus documentos personales "
    "y del vehiculo: licencia de conducir, revision tecnica, seguro de responsabilidad civil, "
    "tarjeta de propiedad y cualquier otro documento exigido por las autoridades venezolanas. "
    "LA PLATAFORMA no asume responsabilidad por infracciones, multas o sanciones derivadas del "
    "incumplimiento de normativas de transito por parte de EL CONDUCTOR.")

clause("NOVENA - PROTECCION DE DATOS",
    "EL CONDUCTOR autoriza a VERONA Ride a almacenar y procesar sus datos personales con la "
    "finalidad exclusiva de operar y mejorar el servicio de la plataforma, en estricto cumplimiento "
    "de las leyes venezolanas aplicables en materia de proteccion de datos personales.")

clause("DECIMA - VIGENCIA Y TERMINACION",
    "Este contrato entra en vigencia a partir de la fecha de firma y tiene duracion indefinida. "
    "Cualquiera de las partes podra darlo por terminado con un aviso previo de 7 (siete) dias "
    "calendario mediante notificacion a traves de la plataforma o correo electronico oficial.")

doc.add_paragraph()

# ── DECLARACION FINAL
p = doc.add_paragraph(
    "EL CONDUCTOR declara que ha leido, comprendido y aceptado voluntariamente el presente "
    "contrato en su totalidad, y procede a estampar su firma a continuacion:")
for r in p.runs:
    font(r, size=10.5)
p.paragraph_format.space_after = Pt(8)

today = datetime.date.today()
p2 = doc.add_paragraph(f"Lugar y fecha: Venezuela, _____ de _______________ de {today.year}")
for r in p2.runs:
    font(r, size=10.5)

doc.add_paragraph()
doc.add_paragraph()

# ── TABLA DE FIRMAS
table = doc.add_table(rows=2, cols=2)
table.style = "Table Grid"

# Fila cabecera
for i, txt in enumerate(["Firma del Conductor", "Firma y Sello VERONA Ride"]):
    cell = table.rows[0].cells[i]
    cell.text = txt
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs:
            font(r, bold=True, size=10.5, color=AZUL)

# Fila de contenido
contenido = [
    "\n\n\n\n_______________________________\nNombre: ________________________\nCedula:  ________________________\nFecha:   ________________________",
    "\n\n\n\n_______________________________\nRepresentante autorizado\nVERONA RIDE VENEZUELA",
]
for i, txt in enumerate(contenido):
    cell = table.rows[1].cells[i]
    cell.text = txt
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs:
            font(r, size=10)

out = "D:/apps/V-Ride_Venezuela/CONTRATO_CONDUCTOR_VERONA_RIDE_VE.docx"
doc.save(out)
print("OK:", out)
